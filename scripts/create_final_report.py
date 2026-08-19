from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('docs/OnTrack_Project_Report_Final.docx')

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, 'FFFFFF')
        shade(table.rows[0].cells[i], '1F4E78')
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table

def bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
        p.add_run(item)

def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(item)

def h(doc, text, level=1):
    doc.add_heading(text, level=level)

def p(doc, text='', bold_prefix=None):
    para = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        para.add_run(bold_prefix).bold = True
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    return para

references = [
'[1] Steel, P. (2007). The nature of procrastination: A meta-analytic and theoretical review of quintessential self-regulatory failure. Psychological Bulletin, 133(1), 65–94. https://doi.org/10.1037/0033-2909.133.1.65',
'[2] Zimmerman, B. J. (1998). Academic studying and the development of personal skill: A self-regulatory perspective. Educational Psychologist, 33(2–3), 73–86. https://doi.org/10.1080/00461520.1998.9653292',
'[3] Gollwitzer, P. M. (1999). Implementation intentions: Strong effects of simple plans. American Psychologist, 54(7), 493–503. https://doi.org/10.1037/0003-066X.54.7.493',
'[4] Locke, E. A., & Latham, G. P. (2002). Building a practically useful theory of goal setting and task motivation. American Psychologist, 57(9), 705–717. https://doi.org/10.1037/0003-066X.57.9.705',
'[5] Baumeister, R. F., Vohs, K. D., & Tice, D. M. (2007). The strength model of self-control. Current Directions in Psychological Science, 16(6), 351–355. https://doi.org/10.1111/j.1467-8721.2007.00534.x',
'[6] Mark, G., Gudith, D., & Klocke, U. (2008). The cost of interrupted work: More speed and stress. Proceedings of CHI 2008, 107–110. https://doi.org/10.1145/1357054.1357072',
'[7] Leroy, S. (2009). Why is it so hard to do my work? The challenge of attention residue when switching between work tasks. Organizational Behavior and Human Decision Processes, 109(2), 168–181. https://doi.org/10.1016/j.obhdp.2009.04.002',
'[8] Monsell, S. (2003). Task switching. Trends in Cognitive Sciences, 7(3), 134–140. https://doi.org/10.1016/S1364-6613(03)00028-7',
'[9] Buehler, R., Griffin, D., & Ross, M. (1994). Exploring the planning fallacy: Why people underestimate their task completion times. Journal of Personality and Social Psychology, 67(3), 366–381. https://doi.org/10.1037/0022-3514.67.3.366',
'[10] Claessens, B. J. C., van Eerde, W., Rutte, C. G., & Roe, R. A. (2007). A review of the time management literature. Personnel Review, 36(2), 255–276. https://doi.org/10.1108/00483480710726136',
'[11] Macan, T. H. (1994). Time management: Test of a process model. Journal of Applied Psychology, 79(3), 381–391. https://doi.org/10.1037/0021-9010.79.3.381',
'[12] Park, S. W., & Sperling, R. A. (2012). Academic procrastinators and their self-regulation. Psychology, 3(12), 12–23. https://doi.org/10.4236/psych.2012.31003',
'[13] van Eerde, W. (2003). Procrastination at work and time management training. The Journal of Psychology, 137(5), 421–434. https://doi.org/10.1080/00223980309600625',
'[14] Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340. https://doi.org/10.2307/249008',
'[15] Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425–478. https://doi.org/10.2307/30036540',
'[16] Fogg, B. J. (2009). A behavior model for persuasive design. Proceedings of Persuasive 2009. https://behaviordesign.stanford.edu/resources/fogg-behavior-model',
'[17] Al-Emran, M., Mezhuyev, V., & Kamalja, M. (2018). Students’ readiness to use mobile technology in higher education. Education and Information Technologies, 23, 109–130. https://doi.org/10.1007/s10639-017-9590-2',
'[18] Crompton, H., Burke, D., & Gregory, K. H. (2017). The use of mobile learning in higher education: A systematic review. Computers & Education, 123, 53–64. https://doi.org/10.1016/j.compedu.2018.04.007',
'[19] Alrasheedi, M., Capretz, L. F., & Raza, A. (2015). A systematic review of the critical success factors for mobile learning. Journal of Computers in Education, 2, 1–15. https://doi.org/10.1007/s40692-014-0028-2',
'[20] Al-Hunaiyyan, A., Alhajri, R., & Alsharhan, S. (2018). Perceptions and challenges of mobile learning in higher education. International Journal of Interactive Mobile Technologies, 12(4), 4–18. https://doi.org/10.3991/ijim.v12i4.8964',
'[21] Oulasvirta, A., Rattenbury, T., Ma, L., & Raita, E. (2012). Habits make smartphone use more pervasive. Personal and Ubiquitous Computing, 16, 105–114. https://doi.org/10.1007/s00779-011-0412-2',
'[22] Pinder, C., Vermeulen, J., Cowan, B. R., & Beale, R. (2018). Digital behaviour change interventions to break and form habits. Interacting with Computers, 30(5), 381–400. https://doi.org/10.1093/iwc/iwy009',
'[23] Kwasnicka, D., Dombrowski, S. U., White, M., & Sniehotta, F. F. (2016). Theoretical explanations for maintenance of behaviour change. Health Psychology Review, 10(3), 277–296. https://doi.org/10.1080/17437199.2016.1151372',
'[24] Design and Implementation of a Time Management Self-Help Mobile App for College Students. (2024). IEEE conference publication. https://ieeexplore.ieee.org/document/10402177',
'[25] M-Learning Excellence: Personalized Mobile Learning for University Students via an Android App. (2024). IEEE conference publication. https://ieeexplore.ieee.org/document/10590121',
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(10.5)
for name, size, color in [('Title', 24, '1F4E78'), ('Heading 1', 16, '1F4E78'), ('Heading 2', 13, '2F75B5'), ('Heading 3', 11, '5B9BD5')]:
    styles[name].font.name = 'Aptos Display'; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color)

# Cover
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run('ONTRACK'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor(31,78,121)
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = para.add_run('A Mobile Deadline and Focus Management Application'); r.bold = True; r.font.size = Pt(18)
doc.add_paragraph()
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.add_run('PROJECT REPORT').bold = True
para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.add_run('Mobile Application, REST API, PostgreSQL Database and Evaluation')
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
for line in ['Course project submission', 'Technology: React Native, Expo, TypeScript, FastAPI and PostgreSQL', 'Repository: https://github.com/thiennguyentuan/OnTrack', 'Prepared: August 2026']:
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER; para.add_run(line)
doc.add_page_break()

h(doc, 'Abstract', 1)
p(doc, 'OnTrack is a mobile application that helps university students turn large deadlines into manageable milestones, tasks and focused work sessions. The application combines planning, execution, post-session review, progress aggregation and deadline risk detection in one workflow. The mobile client is implemented with React Native, Expo and TypeScript. A FastAPI REST backend provides authentication, profile management, planning CRUD, session lifecycle operations, review processing, dashboard data and risk analysis. PostgreSQL stores the user-owned hierarchy and enforces relational integrity. The implemented prototype was validated with automated domain tests, TypeScript checking, REST smoke testing and an end-to-end web flow. This report presents the problem, research background, requirements, analysis and design, implementation, testing, project plan, limitations and future work.')
p(doc, 'Keywords: mobile productivity, time management, academic procrastination, focus session, deadline risk, React Native, FastAPI, PostgreSQL.')

h(doc, 'Table of Contents', 1)
p(doc, 'The following sections are arranged as a submission-ready report. In Microsoft Word, use References > Table of Contents > Update Table if a generated page-number table is required.')

h(doc, '1. Introduction', 1)
h(doc, '1.1 Background and problem', 2)
p(doc, 'University students frequently manage several assignments, projects, examinations and personal commitments at the same time. A calendar entry alone does not explain what should be done today, how much progress is needed, or whether the current pace is sufficient. The problem is therefore not only remembering a date; it is converting a large objective into executable actions and adapting the plan when reality differs from the original estimate.')
p(doc, 'Academic procrastination is commonly discussed as a self-regulation failure [1], while self-regulated learning emphasizes planning, monitoring and reflection [2]. Interruptions and task switching can add cognitive cost and stress [6–8]. These findings motivate an application that combines a small actionable plan, a bounded focus session and a short review loop instead of only displaying a deadline list.')
h(doc, '1.2 Project motivation', 2)
p(doc, 'OnTrack was designed for students who need a lightweight mobile workflow: define a deadline, divide it into milestones and tasks, schedule a session, focus, record the result and see whether the deadline is on track. The core value loop is: Plan → Execute → Review → Track → Adjust.')
h(doc, '1.3 Objectives', 2)
bullets(doc, ['Provide a mobile-first interface for deadlines, milestones, tasks and sessions.', 'Support registration, login, profile update and logout.', 'Allow users to create, read, update and delete planning data in JSON through APIs.', 'Support session scheduling, start, pause, resume, end, review and follow-up.', 'Calculate progress from task to milestone to deadline.', 'Detect deadlines that are at risk or overdue using a transparent rule.', 'Deliver source code and a reproducible PostgreSQL schema for final submission.'])
h(doc, '1.4 Scope', 2)
add_table(doc, ['In scope', 'Out of scope for MVP'], [
('Authentication and profile', 'AI automatic task decomposition'), ('Deadline → milestone → task hierarchy', 'Social network and team collaboration'), ('Focus sessions, timer and review', 'Complex gamification and leaderboard'), ('Progress, history and risk card', 'Two-way Google Calendar sync'), ('Local notification/preferences support', 'Guaranteed blocking of every other app')])

h(doc, '2. Research Background and Design Rationale', 1)
p(doc, 'The literature review was used to translate behavioural and software engineering ideas into concrete features. Goal-setting research supports specific, actionable goals [4]. Implementation intentions support linking a situation to an action, which is reflected in scheduled sessions [3]. Time-management literature emphasizes planning and perceived control [10, 11]. The focus timer and high-focus guidance address interruption and task-switching costs [6–8]. Mobile-learning studies support using a mobile device for contextual, just-in-time interaction, while also warning about usability, privacy and adoption barriers [18–20].')
add_table(doc, ['Research finding', 'OnTrack design response'], [
('Large goals are difficult to self-regulate [1, 2].', 'Deadline → milestone → task decomposition.'),
('Specific goals improve task motivation [4].', 'Task title, priority, target date and position.'),
('If–then plans support goal initiation [3].', 'Planned session with start time and estimated duration.'),
('Interruptions increase stress and switching cost [6–8].', 'Bounded focus timer and High Focus guidance.'),
('People underestimate completion time [9].', 'Actual minutes and post-session review are stored.'),
('Feedback supports self-regulation [2, 10].', 'Review updates progress and enables a follow-up session.'),
('Mobile adoption depends on usefulness and ease of use [14, 15].', 'Three primary tabs: Today, Plans and Me; simple CRUD forms.')])
h(doc, '2.1 Research gap addressed by the project', 2)
p(doc, 'Many productivity tools stop at either a task list, a timer or a calendar. OnTrack connects these parts to a deadline model and calculates a risk signal from actual progress and time remaining. The project is not presented as a clinical intervention or a validated educational experiment; it is a software prototype that operationalizes research-informed principles and creates a platform for later user evaluation.')

h(doc, '3. Requirements Analysis', 1)
h(doc, '3.1 Functional requirements', 2)
add_table(doc, ['ID', 'Requirement', 'Acceptance criterion'], [
('FR-01', 'Register and login', 'Valid users receive a JSON user object and bearer token.'),
('FR-02', 'Profile management', 'Authenticated user can retrieve and update name/timezone.'),
('FR-03', 'Planning CRUD', 'User can create, list, view, edit and delete deadlines, milestones and tasks.'),
('FR-04', 'Session lifecycle', 'A planned session can start, pause, resume and end with state validation.'),
('FR-05', 'Review', 'Ended session review updates session, task, milestone and deadline in one transaction.'),
('FR-06', 'Follow-up', 'Incomplete task can receive a follow-up; completed task cannot.'),
('FR-07', 'Dashboard/history', 'Today, history and risk endpoints return JSON for the mobile screens.'),
('FR-08', 'Database script', 'A fresh PostgreSQL database can be initialized from backend/schema.sql.')])
h(doc, '3.2 Non-functional requirements', 2)
bullets(doc, ['Security: passwords are hashed with bcrypt; API access uses JWT bearer tokens; ownership is checked through the hierarchy.', 'Reliability: review updates related records within one database transaction.', 'Usability: primary navigation is limited to Today, Plans and Me; forms expose loading and error states.', 'Maintainability: feature API modules, domain presentation helpers and backend endpoint groups are separated.', 'Portability: React Native and Expo support Android-first delivery and web verification.', 'Transparency: risk calculation is deterministic and explainable rather than an opaque prediction model.'])

h(doc, '4. Analysis and Design', 1)
h(doc, '4.1 Domain model', 2)
p(doc, 'The central ownership hierarchy is User 1—N Deadline 1—N Milestone 1—N Task 1—N Session. A Task is the work item that can reach 100% completion. A Session is one attempt to work on that task; one task may therefore have multiple sessions. Follow-up sessions retain a previous_session_id so progress history remains traceable.')
add_table(doc, ['Entity', 'Purpose', 'Important attributes'], [
('User', 'Account and ownership boundary', 'id, email, full_name, timezone'), ('Deadline', 'Large objective with due date', 'title, due_at, priority, progress, risk_level'), ('Milestone', 'Intermediate outcome', 'deadline_id, target_at, position, progress'), ('Task', 'Executable work item', 'milestone_id, priority, current_progress, status'), ('Session', 'Bounded work attempt', 'planned_start_at, estimated_minutes, status, review data')])
h(doc, '4.2 Progress and risk rules', 2)
p(doc, 'Task progress is updated from the latest valid review. Milestone progress is the average progress of its tasks, and deadline progress is the average progress of its milestones. The current MVP uses equal weighting because task effort estimation is not yet part of the data model.')
p(doc, 'Expected progress is calculated from time remaining using a simple straight-line pace. The backend derives risk at read time: completed work is ON_TRACK; an unfinished deadline past due is OVERDUE; otherwise a progress gap below -10 percentage points is AT_RISK; all other cases are ON_TRACK. This approach is intentionally understandable and can be replaced by a calibrated model after collecting usage data.')
h(doc, '4.3 System architecture', 2)
p(doc, 'The mobile layer uses Expo Router screens and React Native components. Feature API modules call the REST backend. TanStack Query manages remote data and cache invalidation, while Zustand and AsyncStorage support small local UI/session preferences. The backend is FastAPI with Pydantic validation, psycopg database access, bcrypt password hashing, JWT authentication and SMTP-compatible password reset support. PostgreSQL provides constraints, foreign keys, indexes and cascading deletes.')
add_table(doc, ['Layer', 'Implementation'], [
('Presentation', 'React Native, Expo Router, reusable UI components'), ('Client state', 'TanStack Query for server state; Zustand/AsyncStorage for local state'), ('API boundary', 'Feature API modules calling /api/v3 JSON endpoints'), ('Backend', 'Python FastAPI, Pydantic, JWT, bcrypt'), ('Database', 'PostgreSQL, schema.sql, indexes and foreign keys'), ('Device services', 'Expo notifications, secure storage and Android development build path')])
h(doc, '4.4 Main user flow', 2)
numbered(doc, ['Register or log in.', 'Create a deadline with title, due date and priority.', 'Open the deadline and add milestones.', 'Add tasks under a milestone.', 'Plan a focus session for a task.', 'Start the session, optionally pause/resume, then end it.', 'Submit a review with actual progress and a result note.', 'The system updates task, milestone and deadline progress.', 'If progress is incomplete, create a follow-up session; if the deadline is behind schedule, display the risk card.'])

h(doc, '5. Implementation', 1)
h(doc, '5.1 Mobile application', 2)
p(doc, 'The mobile source is in the repository root. The route structure includes authentication screens, Today/Plans/Me tabs, deadline/milestone/task detail and CRUD screens, session planning, focus and review screens, history, progress, risk and settings. The app is Android-first but can be exported to web for development and demonstration.')
h(doc, '5.2 API/backend', 2)
p(doc, 'The backend source is under backend/app/main.py. Implemented endpoint groups include /health, /api/v3/auth/*, /api/v3/users/me, /api/v3/deadlines, /api/v3/milestones, /api/v3/tasks, /api/v3/sessions/*, /api/v3/dashboard/today and deadline risk. Responses are JSON and authenticated routes require Authorization: Bearer <token>.')
h(doc, '5.3 Database', 2)
p(doc, 'The standalone PostgreSQL initialization script is backend/schema.sql. It creates the users, deadlines, milestones, tasks, sessions and password_reset_tokens tables, constraints, cascading relationships and indexes. The repository also contains Supabase migrations and supabase/seed.sql for the alternative database path documented in the project architecture.')
h(doc, '5.4 Submission links and reproducibility', 2)
add_table(doc, ['Deliverable', 'Location'], [
('GitHub repository', 'https://github.com/thiennguyentuan/OnTrack'),
('Mobile application source', 'Repository root: app/, src/, components/, assets/, package.json'),
('API/backend source', 'Repository: backend/app/main.py and backend/requirements.txt'),
('Database script', 'Repository: backend/schema.sql'),
('Demo seed data', 'Repository: backend/seed_demo.py and supabase/seed.sql'),
('Existing project documents', 'Repository: OnTrack_Project_Documentation/ and docs/')])

h(doc, '6. Project Plan', 1)
p(doc, 'The project was organized as six implementation sprints. The plan follows the three-member course structure: mobile frontend, backend/database and testing/documentation/integration. All members should participate in analysis, design, implementation, testing, documentation and presentation, even when each person owns a primary area.')
add_table(doc, ['Sprint', 'Activities', 'Outputs', 'Verification'], [
('1. Discovery', 'Problem definition, target users, scope, use cases and backlog', 'Requirements, use cases, initial backlog', 'Requirement review'),
('2. Design', 'Navigation, wireframes, domain model, ERD and architecture', 'Wireframes, ERD, technical architecture', 'Design walkthrough'),
('3. Foundation', 'Expo project, routing, theme, FastAPI service and PostgreSQL schema', 'Running mobile shell, API health, schema script', 'Build and health check'),
('4. Planning features', 'Auth, profile, deadline, milestone and task CRUD', 'Plans screens and REST resources', 'CRUD tests and ownership checks'),
('5. Focus loop', 'Session planning, timer, lifecycle, review, follow-up and dashboard', 'Today, Focus, Review, History and risk features', 'REST smoke and domain tests'),
('6. Integration and submission', 'E2E flow, bug fixing, documentation, report and presentation preparation', 'Final repository, report, database script, test evidence', 'Typecheck, tests, export and demo')])
h(doc, '6.1 Risk management plan', 2)
add_table(doc, ['Risk', 'Impact', 'Mitigation'], [
('Mobile/backend API mismatch', 'High', 'Keep feature API modules and JSON contracts; use integration smoke flow.'),
('Timer becomes inaccurate in background', 'High', 'Use timestamps expected_end_at instead of decrementing a counter.'),
('User data leakage', 'High', 'Require JWT and verify ownership through joins before CRUD.'),
('Database unavailable during demo', 'High', 'Provide Docker Compose and backend/schema.sql; document startup steps.'),
('High Focus cannot control all notifications', 'Medium', 'Explain Android permission and use development build for native testing.'),
('Scope expansion', 'Medium', 'Keep AI, collaboration and calendar sync out of MVP.')])

h(doc, '7. Testing and Results', 1)
h(doc, '7.1 Test strategy', 2)
bullets(doc, ['Unit/domain tests cover presentation mapping, progress rules, navigation, settings, password reset and session timer behaviour.', 'Type checking uses pnpm run typecheck.', 'REST integration tests cover authentication, planning and session flows where the local backend/database is available.', 'Smoke testing follows register → deadline → milestone → task → session → start → end → review → follow-up.', 'Web export provides an additional reproducible demonstration target for the Expo application.'])
h(doc, '7.2 Verification evidence', 2)
add_table(doc, ['Check', 'Recorded result'], [
('pnpm run typecheck', 'PASS in the project integration report.'), ('pnpm test', 'PASS in the project integration report; integration suites requiring unavailable Supabase were documented separately.'), ('git diff --check', 'PASS in the project integration report.'), ('GET /health', 'Returns status=ok and database=postgresql when the local service is running.'), ('REST smoke flow', 'Register → planning hierarchy → session lifecycle → review → follow-up verified in the project report.'), ('Expo web export', 'Bundle completed successfully in the documented integration run.')])
p(doc, 'The current report intentionally distinguishes recorded verification from a new run. Before submission, run the commands in Section 9 on the final machine and attach the latest output if the instructor requires dated evidence.')
h(doc, '7.3 Known limitations', 2)
bullets(doc, ['The MVP risk model is a heuristic and has not been validated against a longitudinal student dataset.', 'Progress uses equal averages and does not weight task effort or complexity.', 'Focus mode can guide or open Android Do Not Disturb settings but cannot guarantee blocking every notification.', 'Some preferences are local to the device because the current REST schema does not persist all preference fields.', 'The project is Android-first; a complete production iOS release is future work.', 'The repository contains both the FastAPI/PostgreSQL implementation and historical Supabase documentation; the FastAPI/PostgreSQL path is the current demo backend.'])

h(doc, '8. Future Work', 1)
numbered(doc, ['Run a user study with students using pre/post measures for planning behaviour, focus-session completion and perceived workload.', 'Replace equal-average progress with effort-weighted progress using estimated minutes or task weights.', 'Learn personal completion pace from historical sessions and calibrate risk thresholds per user.', 'Add server-side user settings and notification synchronization across devices.', 'Implement robust offline retry and conflict handling for active sessions.', 'Add native Android notification policy access and test High Focus in an Expo development build.', 'Add accessibility audit, localization and production error monitoring.', 'Evaluate calendar integration, but keep it optional and privacy-preserving.', 'Introduce secure deployment configuration, secret management, rate limiting and refresh-token rotation before production use.'])

h(doc, '9. How to Run and Demonstrate', 1)
numbered(doc, ['Install Node.js dependencies from the repository root with pnpm install.', 'Start PostgreSQL from the repository root with docker compose up -d postgres.', 'Install backend dependencies in backend/.venv using backend/requirements.txt.', 'Start FastAPI with uvicorn app.main:app --reload --port 8000 from backend.', 'Set EXPO_PUBLIC_API_URL to the backend address; a physical phone must use the computer LAN IP instead of localhost.', 'Start Expo with pnpm start, then demonstrate register, create planning hierarchy, plan and run a session, review it and open the risk/history screens.', 'Run pnpm run typecheck and pnpm run test before final submission.'])
p(doc, 'The backend Swagger documentation is available at http://localhost:8000/docs while the local API is running. Do not commit backend/.env, passwords, tokens or other secrets.')

h(doc, '10. Conclusion', 1)
p(doc, 'OnTrack satisfies the core course deliverables: a mobile application source codebase, a JSON REST API/backend, a PostgreSQL database script and a documented project plan. Its main contribution is the integration of deadline decomposition, focused execution, review-based progress updates and transparent risk feedback. The implemented prototype is suitable for demonstration and further evaluation. Future work should focus on empirical user validation, stronger security/deployment practices, effort-weighted planning and native notification controls.')

h(doc, 'References: 25 Research Papers and Supporting Sources', 1)
p(doc, 'The following references were selected for the project background. DOI and publisher links were checked against scholarly search results and publisher/index pages during report preparation. References [1]–[23] provide theoretical and empirical foundations; [24]–[25] are closely related mobile-app implementations.')
for ref in references:
    para = doc.add_paragraph(style='List Number')
    para.add_run(ref)

h(doc, 'Appendix A. API Endpoint Summary', 1)
add_table(doc, ['Method', 'Endpoint', 'Purpose'], [
('GET', '/health', 'Database/service health check'), ('POST', '/api/v3/auth/register', 'Register'), ('POST', '/api/v3/auth/login', 'Login'), ('GET/PUT', '/api/v3/users/me', 'Read/update profile'), ('GET/POST', '/api/v3/deadlines', 'List/create deadlines'), ('GET/PUT/DELETE', '/api/v3/deadlines/{id}', 'Deadline detail and CRUD'), ('GET/POST', '/api/v3/milestones', 'List/create milestones'), ('GET/POST', '/api/v3/tasks', 'List/create tasks'), ('POST', '/api/v3/sessions', 'Create session'), ('POST', '/api/v3/sessions/{id}/start|pause|resume|end', 'Session lifecycle'), ('POST', '/api/v3/sessions/{id}/review', 'Persist review and progress'), ('GET', '/api/v3/dashboard/today', 'Today dashboard'), ('GET', '/api/v3/sessions/history', 'Session history'), ('GET', '/api/v3/deadlines/{id}/risk', 'Risk analysis')])

h(doc, 'Appendix B. Database Tables', 1)
add_table(doc, ['Table', 'Primary role', 'Relationship'], [
('users', 'Account and owner', 'One user owns many deadlines'), ('deadlines', 'Large objective', 'One deadline contains many milestones'), ('milestones', 'Intermediate outcome', 'One milestone contains many tasks'), ('tasks', 'Executable work', 'One task has many sessions'), ('sessions', 'Focus attempt and review', 'May reference a previous session'), ('password_reset_tokens', 'One-time reset token', 'Belongs to a user')])

# Footer
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('OnTrack Project Report | ')
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); footer._p.append(fld)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
print('references', len(references))
